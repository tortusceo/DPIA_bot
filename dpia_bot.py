#import relevant libraries
#word or excel to markdown converter, google generative - call gemini-2.5-pro, markdown to word/excel converter, download file

import os
from dotenv import load_dotenv
import google.generativeai as genai
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH # For heading check, if used
import pandas as pd # Renamed for convention
import pypandoc # Added for Markdown to DOCX conversion

# Global constant for the reference DPIA file
REFERENCE_DPIA_FILENAME = "reference_dpia.txt"
# Placeholder for customer uploaded file (for testing)
CUSTOMER_FILE_PLACEHOLDER = "customer_template.docx" # Or .xlsx
# PANDOC_REFERENCE_DOCX is removed as we'll use the customer's file directly

def load_google_api_key():
    """Loads the Google API key from .env file."""
    load_dotenv()
    api_key = os.getenv("GOOGLE_API_KEY")
    if not api_key:
        raise ValueError("GOOGLE_API_KEY not found in .env file or environment variables.")
    return api_key

def extract_text(filepath: str) -> str:
    """Extracts all text from a given text file (UTF-8 encoded)."""
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            return f.read()
    except FileNotFoundError:
        print(f"Error: File not found at {filepath}")
        return ""
    except Exception as e:
        print(f"Error reading file {filepath}: {e}")
        return ""

def upload_file_placeholder(original_filename: str) -> str:
    """
    Placeholder for file upload. In a real app, this would handle the upload process.
    For now, it just returns the provided filename, assuming it exists locally for processing.
    """
    print(f"Simulating upload of file: {original_filename}")
    if not os.path.exists(original_filename):
        print(f"Warning: Placeholder upload file '{original_filename}' does not exist.")
    return original_filename

def get_file_extension(filepath: str) -> str:
    """Extracts the file extension from a filepath."""
    return os.path.splitext(filepath)[1].lower()

def convert_file_to_markdown(filepath: str) -> str:
    """
    Converts the given file (Word, Excel) to Markdown text.
    """
    file_extension = get_file_extension(filepath)
    print(f"Attempting to convert {filepath} (type: {file_extension}) to Markdown...")
    markdown_parts = []

    if file_extension == '.docx':
        try:
            doc = docx.Document(filepath)
            for para in doc.paragraphs:
                # Basic heading detection (can be improved by checking styles more robustly)
                if para.style.name.startswith('Heading'):
                    level = 1
                    try:
                        level = int(para.style.name[-1])
                        if not (1 <= level <= 6): level = 1
                    except ValueError:
                        level = 1 # Default to H1 if style name isn't like 'Heading 1'
                    markdown_parts.append("#" * level + " " + para.text)
                else:
                    markdown_parts.append(para.text)
            
            if doc.tables:
                markdown_parts.append("\n\n--- Tables ---")
                for i, table in enumerate(doc.tables):
                    markdown_parts.append(f"\n**Table {i+1}**")
                    header = [cell.text for cell in table.rows[0].cells]
                    markdown_parts.append("| " + " | ".join(header) + " |")
                    markdown_parts.append("|" + " --- |" * len(header))
                    for row in table.rows[1:]:
                        row_content = [cell.text for cell in row.cells]
                        markdown_parts.append("| " + " | ".join(row_content) + " |")
                    markdown_parts.append("") # Add a blank line after table
            
            print("DOCX to Markdown conversion completed.")
            return "\n\n".join(markdown_parts) # Separate paragraphs/elements with double newline
        except Exception as e:
            print(f"Error converting DOCX {filepath} to Markdown: {e}")
            return "Error during DOCX conversion."
            
    elif file_extension in ['.xlsx', '.xls']:
        try:
            xls = pd.ExcelFile(filepath)
            for sheet_name in xls.sheet_names:
                markdown_parts.append(f"## Sheet: {sheet_name}\n")
                df = xls.parse(sheet_name)
                # Convert dataframe to markdown table, handling potential NaNs
                df_filled = df.fillna('') # Replace NaN with empty string for cleaner markdown
                markdown_parts.append(df_filled.to_markdown(index=False))
                markdown_parts.append("\n")
            print("Excel to Markdown conversion completed.")
            return "\n".join(markdown_parts)
        except Exception as e:
            print(f"Error converting Excel {filepath} to Markdown: {e}")
            return "Error during Excel conversion."
            
    elif file_extension in ['.md', '.txt']:
        print("File is already Markdown/text. Reading directly.")
        return extract_text(filepath)
    else:
        print(f"Unsupported file type for conversion: {file_extension}")
        return ""

def generate_dpia_from_prompt(full_prompt_str: str, api_key: str) -> str:
    """
    Uses Google Gemini API to generate the DPIA content based on the full prompt.
    """
    print("Calling Google Gemini API... this can take a few minutes.")
    try:
        genai.configure(api_key=api_key)
        # Note: Using 'gemini-pro' as a common model. 
        # The prompt mentions 'gemini-2.5-pro', ensure this model name is correct for API usage if needed.
        model = genai.GenerativeModel('gemini-2.5-pro-preview-05-06') 
        
        # Safety settings can be configured here if needed, for example:
        # from google.generativeai.types import HarmCategory, HarmBlockThreshold
        # safety_settings = {
        #     HarmCategory.HARM_CATEGORY_HARASSMENT: HarmBlockThreshold.BLOCK_NONE,
        #     HarmCategory.HARM_CATEGORY_HATE_SPEECH: HarmBlockThreshold.BLOCK_NONE,
        #     HarmCategory.HARM_CATEGORY_SEXUALLY_EXPLICIT: HarmBlockThreshold.BLOCK_NONE,
        #     HarmCategory.HARM_CATEGORY_DANGEROUS_CONTENT: HarmBlockThreshold.BLOCK_NONE,
        # }
        # response = model.generate_content(full_prompt_str, safety_settings=safety_settings)

        response = model.generate_content(full_prompt_str)
        
        # Debugging: print parts of the response if text is empty
        if not response.text:
            print("Warning: Gemini API returned an empty text response.")
            print(f"Prompt Parts: {response.prompt_feedback}")
            # Consider how to handle blocked prompts or other issues.
            # For now, return an error message.
            return "Error: Gemini API returned no text. The prompt might have been blocked or an issue occurred."
            
        print("Successfully received response from Gemini API.")
        return response.text
    except Exception as e:
        print(f"Error calling Google Gemini API: {e}")
        # Provide more context if possible, e.g., if it's an authentication error vs. model error
        return f"Error during Gemini API call: {str(e)}"

def convert_markdown_to_original_format(
    markdown_content: str, 
    original_file_extension: str, 
    output_filename: str,
    style_reference_filepath: str | None = None # Added to use original customer file for styling
) -> str:
    """
    Converts the generated Markdown back to the original file format (e.g., DOCX, basic Excel).
    Uses Pandoc via pypandoc for .docx output for better formatting.
    If style_reference_filepath is provided (and is .docx), it's used as Pandoc's reference-doc.
    """
    print(f"Attempting to convert Markdown back to {original_file_extension} format...")
    try:
        if original_file_extension == '.docx':
            try:
                pandoc_args = []
                # Use the provided style_reference_filepath if it's a docx and exists
                if style_reference_filepath and get_file_extension(style_reference_filepath) == '.docx' and os.path.exists(style_reference_filepath):
                    pandoc_args.append(f'--reference-doc={style_reference_filepath}')
                    print(f"Using Pandoc reference DOCX: {style_reference_filepath}")
                else:
                    print("Style reference DOCX not provided, not found, or not a .docx file. Using Pandoc default styles.")

                # Temporarily print the markdown content for table debugging
                print(f"--- MARKDOWN INPUT TO PANDOC ---")
                # Print a snippet if it's too long, or indicate length
                if len(markdown_content) > 2000:
                    print(markdown_content[:1000] + "... [TRUNCATED] ..." + markdown_content[-1000:])
                    print(f"(Total length: {len(markdown_content)})")
                else:
                    print(markdown_content)
                print(f"--- END MARKDOWN INPUT TO PANDOC ---")

                pypandoc.convert_text(
                    markdown_content, 
                    'docx', 
                    format='md', 
                    outputfile=output_filename,
                    extra_args=pandoc_args
                )
                print(f"Markdown converted and saved to DOCX using Pandoc: {output_filename}")
                return output_filename
            except Exception as e:
                print(f"Error using Pandoc for DOCX conversion: {e}")
                print("Pandoc might not be installed or there was an issue with the conversion.")
                print("Falling back to basic python-docx conversion (formatting will be limited).")
                doc = docx.Document()
                for para_text in markdown_content.split('\n\n'):
                    if para_text.strip():
                        doc.add_paragraph(para_text)
                doc.save(output_filename)
                print(f"Basic Markdown conversion saved to DOCX: {output_filename}")
                return output_filename
            
        elif original_file_extension in ['.xlsx', '.xls']:
            # Simplified: Save Markdown content to a text file, but name it with .xlsx.
            # A true Markdown to Excel conversion is complex and would require
            # parsing Markdown tables and structuring them into sheets.
            # For a slightly better approach, if markdown_content is CSV-like, pandas could parse and save.
            # For now, just saving the markdown text directly into a file with excel extension.
            try:
                # Attempt to write to an Excel file using pandas if content is simple (e.g. table-like)
                # This is a basic attempt, might fail for complex markdown.
                # A robust solution would involve parsing markdown tables into DataFrames.
                # For now, we'll just write the raw markdown to a sheet if it's not easily parsable as a table.
                
                # Quick check if it's a markdown table (very naive)
                if "|---|---|:" in markdown_content or ("|" in markdown_content and "\n" in markdown_content):
                    try:
                        # Try to read markdown tables if any (requires 'tabulate' often, which is not listed)
                        # This is a placeholder for a more robust Markdown table to Excel conversion
                        # For now, writing the whole content to a single cell might be safer.
                        df = pd.DataFrame([x.split('|') for x in markdown_content.strip().split('\n') if x.strip()])
                        # This parsing is very naive.
                        with pd.ExcelWriter(output_filename, engine='openpyxl') as writer:
                            df.to_excel(writer, sheet_name='Sheet1', index=False, header=False)
                        print(f"Markdown (table attempt) converted and saved to Excel: {output_filename}")
                    except Exception as table_parse_error:
                        print(f"Could not parse Markdown as table for Excel, saving as text: {table_parse_error}")
                        # Fallback: save markdown content into a text cell
                        df_fallback = pd.DataFrame([["Markdown Content"], [markdown_content]])
                        with pd.ExcelWriter(output_filename, engine='openpyxl') as writer:
                            df_fallback.to_excel(writer, sheet_name='Markdown_Output', index=False, header=False)
                        print(f"Markdown content saved to an Excel sheet (text): {output_filename}")
                else:
                     # Fallback: save markdown content into a text cell
                    df_fallback = pd.DataFrame([["Markdown Content"], [markdown_content]])
                    with pd.ExcelWriter(output_filename, engine='openpyxl') as writer:
                        df_fallback.to_excel(writer, sheet_name='Markdown_Output', index=False, header=False)
                    print(f"Markdown content saved to an Excel sheet (text): {output_filename}")

            except Exception as pd_e:
                print(f"Error saving Markdown to Excel using pandas: {pd_e}. Saving as plain text with .xlsx extension.")
                # Fallback to plain text if pandas method fails
                with open(output_filename, 'w', encoding='utf-8') as f:
                    f.write("This Excel file contains Markdown content that could not be fully converted to structured Excel data.\n\n")
                    f.write(markdown_content)
                print(f"Markdown saved as text content within {output_filename} (simulating Excel).")
            return output_filename
            
        else: # Assuming .md or .txt
            with open(output_filename, 'w', encoding='utf-8') as f:
                f.write(markdown_content)
            print(f"Saved Markdown directly to {output_filename}.")
            return output_filename
    except Exception as e:
        print(f"Error in convert_markdown_to_original_format for {output_filename}: {e}")
        # Fallback: Save the markdown content to a .txt file to avoid data loss
        txt_fallback_filename = os.path.splitext(output_filename)[0] + "_conversion_error.txt"
        with open(txt_fallback_filename, 'w', encoding='utf-8') as f:
            f.write(markdown_content)
        print(f"Content saved to {txt_fallback_filename} due to conversion error.")
        return txt_fallback_filename

def download_file_placeholder(output_filepath: str):
    """
    Placeholder for allowing the user to download the file.
    In a web app, this would trigger a browser download.
    For a script, it just confirms the file is ready.
    """
    if os.path.exists(output_filepath):
        print(f"File '{output_filepath}' is ready for 'download' (available at this path).")
    else:
        print(f"Error: Output file '{output_filepath}' not found for download.")

# New function to construct the prompt
def construct_prompt_for_gemini(reference_dpia_text: str, customer_template_markdown: str) -> str:
    """Constructs the full prompt string for the Gemini API."""
    return f""" 
<scenario> You are a highly skilled, medical device and data governance specialist. Your responses are audited and any mistakes will incur
large financial penalties to the company, you must be 100% accurate. <scenario>

<task> You are given a reference Data Protection Impact Assessment text (below) and the user has uploaded a file from a customer asking to fill 
out their DPIA template. You need to transfer the information from the reference into the new format. <task>

<output> The output should be a complete markdown file preserving the original formatting of the doc uploaded completed and filled out with the DPIA information. 
If there are areas where the reference text does NOT contain sufficient information to answer, list these in a seperate part of the output called clarifications in tags like so:
<clarification> 
1. The customer DPIA has asked for a Disaster Recovery Plan at section [x.1] - this is not available in the reference DPIA. Please address
</clarification>

<constraints>
The completed DPIA should be as thorough as possible given the reference text, but MUST NOT contain any information that isn't directly entailed from the reference. Preferably use VERBATIM SENTENCES from the reference to complete the customer text, remembering this is a Class 1 medical device and data security and claims are of the upmost importance to be 100% accurate at all times. If there is insufficient information use the clarification system where needed. Partial answers are permissible but leave a tag [partial - to be completed - see clarification X] at the end of answers which need further clarification to be filled out later by a human reviewer.
<constraints>

<additional_context>
TORTUS is a Class 1 MHRA registered medical device that takes data security, governance and safety incredibly seriously. The work we do is using AI in real-world hospitals and clinics, where life and death and highly sensitive information is shared on a daily basis, so the system must always be satte of the art and of clinical-grade quality at ALL TIMES. This task is to help our compliance team with external requests for information about our data protection information documents (DPIAs) - these are a legal requirement under GDPR but often come in a different format to the one we supply. While technically our document (the reference) is complete, to help enterprise sales and commercial we offer a service to help the enterprise fill out their details and close deals faster. So this task is critically important from both a commercial and compliance perspective. The full pipeline is: customer sends us a new DPIA, we pass this through your service first and produces a first-pass automated DPIA in the format required (CSV/DOCX) (handled in code after this prompt), this is reviewed by a human compliance support officer, and then returned to the customer, saving 50-80% of time of human-in-the-loop work and enabling the companies overall goals. Lastly our company values are be kind (Act with empathy), stay curious (back everything with data), patient first (always consider the patient) and kick ass (aim for excellence). 
</additional_context>

<reference_text>
This is our entire markdown formatted DPIA - use this as the reference: {reference_dpia_text}
</reference_text>

<customer_DPIA_template>
This is the customer DPIA we have recieved to fill out - it is a markdown conversion of either a docx or CSV file and needs to be returnned later in the pipeline in the same format, so preserve the markdown format while answering to enable this.
{customer_template_markdown}
</customer_DPIA_template>
 """

def main():
    """Main function to run the DPIA generation bot (CLI version)."""
    try:
        api_key = load_google_api_key()
        print("Google API Key loaded successfully.")
    except ValueError as e:
        print(e)
        return

    print(f"Loading reference DPIA: {REFERENCE_DPIA_FILENAME}")
    extract = extract_text(REFERENCE_DPIA_FILENAME)
    if not extract:
        print(f"Could not load reference DPIA content from {REFERENCE_DPIA_FILENAME}. Exiting.")
        return

    customer_filepath_to_process = upload_file_placeholder(CUSTOMER_FILE_PLACEHOLDER)
    if not os.path.exists(customer_filepath_to_process):
        print(f"Customer file '{customer_filepath_to_process}' not found. Please create it or change CUSTOMER_FILE_PLACEHOLDER. Exiting.")
        return

    customer_original_file_extension = get_file_extension(customer_filepath_to_process)
    print(f"Customer file type detected as: {customer_original_file_extension}")

    conversion = convert_file_to_markdown(customer_filepath_to_process)
    if not conversion or "Error during" in conversion:
        print(f"Could not convert customer file '{customer_filepath_to_process}' to Markdown. Exiting.")
        return
    print(f"Customer file converted to Markdown successfully. Length: {len(conversion)}")

    # Use the new function to construct the prompt
    prompt_text_for_gemini = construct_prompt_for_gemini(extract, conversion)
    print("Prompt for Gemini constructed.")

    new_dpia_markdown = generate_dpia_from_prompt(prompt_text_for_gemini, api_key)
    if "Error during Gemini API call" in new_dpia_markdown or "Error: Gemini API returned no text" in new_dpia_markdown:
        print(f"Failed to generate DPIA from Gemini: {new_dpia_markdown}")
        print("Exiting due to Gemini API error.")
        return
    print("Received generated DPIA from Gemini.")

    output_filename_base = os.path.splitext(os.path.basename(customer_filepath_to_process))[0]
    prepared_doc_filepath = f"{output_filename_base}_completed{customer_original_file_extension}"
    
    final_output_path = convert_markdown_to_original_format(
        new_dpia_markdown, 
        customer_original_file_extension, 
        prepared_doc_filepath,
        style_reference_filepath=customer_filepath_to_process
    )

    if "_conversion_error.txt" in final_output_path:
        print(f"There was an error converting Markdown to the original format. Result saved in {final_output_path}")
    else:
        print(f"Converted new DPIA to original format: {final_output_path}")

    download_file_placeholder(final_output_path)

if __name__ == "__main__":
    # Create dummy files for testing if they don't exist
    if not os.path.exists(REFERENCE_DPIA_FILENAME):
        with open(REFERENCE_DPIA_FILENAME, "w", encoding="utf-8") as f:
            f.write("# Reference DPIA Content\\n\\nThis is sample markdown for the reference DPIA.")
        print(f"Created dummy '{REFERENCE_DPIA_FILENAME}'")

    if not os.path.exists(CUSTOMER_FILE_PLACEHOLDER):
        try:
            doc = docx.Document()
            doc.add_heading('Customer DPIA Template', 0)
            doc.add_paragraph('Section 1: What is your name?')
            doc.add_paragraph('Section 2: What is your quest?')
            # Add a simple table for testing table conversion
            table = doc.add_table(rows=2, cols=2)
            table.cell(0,0).text = "Header A"
            table.cell(0,1).text = "Header B"
            table.cell(1,0).text = "Data 1"
            table.cell(1,1).text = "Data 2"
            doc.save(CUSTOMER_FILE_PLACEHOLDER)
            print(f"Created dummy '{CUSTOMER_FILE_PLACEHOLDER}' with a table.")
        except Exception as e:
            print(f"Could not create dummy DOCX {CUSTOMER_FILE_PLACEHOLDER} (ensure python-docx is installed): {e}")
            print(f"Please create a dummy '{CUSTOMER_FILE_PLACEHOLDER}' manually (e.g., a .docx or .xlsx file).")
            
    if not os.path.exists(".env"):
        with open(".env", "w", encoding="utf-8") as f:
            f.write('GOOGLE_API_KEY="YOUR_API_KEY_HERE"') # Ensure quotes are correct for .env
        print("Created dummy '.env' file. Please add your GOOGLE_API_KEY to it.")

    main()